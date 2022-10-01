import requests
from bs4 import BeautifulSoup
from datetime import date
from docx import Document
import pandas as pd

today = date.today()

page = requests.get("https://inshorts.com/en/read/")
page.status_code
# A status_code of 200 means that the page downloaded successfully.
soup = BeautifulSoup(page.content, 'html.parser')

headlines_tags = soup.find_all('span', itemprop="headline")
headlines = [pt.get_text() for pt in headlines_tags]
author_tags = soup.find_all('span', class_="author")
author = [pt.get_text() for pt in author_tags]
author = author[::2]
date_tags = soup.find_all('span', class_="date")
date = [pt.get_text() for pt in date_tags]
time_tags = soup.find_all('span', class_="time", itemprop="datePublished")
time = [pt.get_text() for pt in time_tags]
body_tags = soup.find_all('div', itemprop="articleBody")
body = [pt.get_text() for pt in body_tags]

final = []
for i in range(len(body)):
    news = {'Title': None,
            'Author': None,
            'Date': None,
            'Time': None,
            'Details': None}
    news['Title'] = headlines[i]
    news['Author'] = author[i]
    news['Date'] = date[i]
    news['Time'] = time[i]
    news['Details'] = body[i]
    final.append(news)
final

textfile = open("news.txt", "w", encoding="utf-8")
textfile.write("NEWS " + str(today) + "\n")
for element in final:
    for key, value in element.items():
        textfile.write(f'{key}: {value}\n')
    textfile.write('\n')
textfile.close()

document = Document()
document.add_heading('NEWS ' + str(today), 0)
for element in final:
    p = document.add_paragraph('')
    for key, value in element.items():
        p.add_run(f'{key}: {value}\n')
    p.add_run('\n')
document.save('D:/pythonProject2/news.docx')

news_df = pd.DataFrame({
    "Title": headlines,
    "Author": author,
    "Date": date,
    "Time": time,
    "Details": body
})
news_df.to_csv('news.csv')
