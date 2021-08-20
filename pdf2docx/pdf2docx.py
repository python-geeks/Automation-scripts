# Import Libraries
from pdf2docx import parse
from typing import Tuple
import groupdocs_conversion_cloud
from PIL import Image
import pytesseract
import sys
from pdf2image import convert_from_path
import os
from docx import *

PDF_file = "C:/Users/pankaj/Desktop/RESUME1.pdf"
  
'''
Part #1 : Converting PDF to images
So here as a example i have included my resume and management file
'''
def pdf_to_image(PDF_file):  
    # Store all the pages of the PDF in a variable
    pages = convert_from_path(PDF_file, 500)
  
    # Counter to store images of each page of PDF to image
    image_counter = 1
  
    # Iterate through all the pages stored above
    for page in pages:
  
        # Declaring filename for each page of PDF as JPG
        # For each page, filename will be:
        # PDF page 1 -> page_1.jpg
        # PDF page 2 -> page_2.jpg
        # PDF page 3 -> page_3.jpg
        # ....
        # PDF page n -> page_n.jpg
        filename = "page_"+str(image_counter)+".jpg"
      
        # Save the image of the page in system
        page.save(filename, 'JPEG')
  
        # Increment the counter to update filename
        image_counter = image_counter + 1
    return image_counter
  
'''
Part #2 - Recognizing text from the images using OCR
'''
def image_to_docx_ocr(image_counter , save_file = 'out.docx'):
    # Variable to get count of total number of pages
    filelimit = image_counter-1
  
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    # Iterate from 1 to total number of pages
    for i in range(1, filelimit + 1):
      
        # Set filename to recognize text from
        # Again, these files will be:
        # page_1.jpg
        # page_2.jpg
        # ....
        # page_n.jpg
        filename = "page_"+str(i)+".jpg"
          
        # Recognize the text as string in image using pytesserct
        text = str(((pytesseract.image_to_string(Image.open(filename)))))
  
        # The recognized text is stored in variable text
        # Any string processing may be applied on text
        # Here, basic formatting has been done:


        document = Document()
        for i in text.split("\n\n"):
            try:
                document.add_paragraph(i)
            except:
                i = i.replace(""," ")
                try:
                    document.add_paragraph(i)
                except:
                    print("Character Not Recognised as ASCII value : " , i)
    # Finally, we save the document .
    document.save(save_file)

def convert_pdf2docx(input_file: str, pages: Tuple = None):

    if pages:
        pages = [int(i) for i in list(pages) if i.isnumeric()]
    result = parse(pdf_file=input_file, pages=pages)
    print("###### Conversion Complete #######")
    return result


def cloud_convert(app_sid, app_key):
    convert_api = groupdocs_conversion_cloud.ConvertApi.from_keys(app_sid, app_key)     # Create instance of the API
    file_api = groupdocs_conversion_cloud.FileApi.from_keys(app_sid, app_key)

    try:
        # upload source file to storage
        filename = input("Please enter the complete and correct file path of the required pdf: ")
        remote_name = filename
        output_name = remote_name.split('.')[0] + "docx"
        remote_name = filename
        strformat = 'docx'

        request_upload = groupdocs_conversion_cloud.UploadFileRequest(remote_name, filename)
        file_api.upload_file(request_upload)

# Extract Text from PDF document
        settings = groupdocs_conversion_cloud.ConvertSettings()
        settings.file_path = remote_name
        settings.format = strformat
        settings.output_path = output_name
        request = groupdocs_conversion_cloud.ConvertDocumentRequest(settings)
        response = convert_api.convert_document(request)
        print("Document converted successfully: " + str(response))

    except groupdocs_conversion_cloud.ApiException as e:
        print("Exception when calling get_supported_conversion_types: {0}".format(e.message))


if __name__ == "__main__":

    choice = input("Press 1 for basic file conversion, 2 for advanced file conversion(GroupDocs account required), 3 for the OCR over scanned pdf!")
    while choice != '1' and choice != '2':
        print("Please enter a valid choice!\n")
        choice = input("Press 1 for basic file conversion, 2 for advanced file conversion(GroupDocs account required)")
    if choice == '1':
        pdf_path = input("Please enter the complete and correct file path of the required pdf: ")
        convert_pdf2docx(pdf_path)
    elif choice == '2':
        sid = input("Enter app SID: ")
        api_key = input("Enter api key: ")
        cloud_convert(sid, api_key)
    else:
        image_counter = pdf_to_image(PDF_file)
        image_to_docx_ocr(image_counter = image_counter )
        print('done')
